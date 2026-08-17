"""
Đọc nội dung từ file (txt/docx/pdf), dịch, và ghi ra file kết quả
cùng định dạng (docx/pdf giữ cấu trúc đoạn văn cơ bản; pdf tạo lại dạng văn bản thuần).
"""
import os
from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT

from .translator import translate_batch


def _ext(path: str) -> str:
    return os.path.splitext(path)[1].lower()


def extract_paragraphs(path: str) -> list[str]:
    """Trích các đoạn văn bản khác rỗng từ file, tuỳ theo định dạng."""
    ext = _ext(path)

    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return [line.rstrip("\n") for line in f]

    if ext == ".docx":
        doc = Document(path)
        return [p.text for p in doc.paragraphs]

    if ext == ".pdf":
        reader = PdfReader(path)
        paragraphs = []
        for page in reader.pages:
            text = page.extract_text() or ""
            paragraphs.extend(text.split("\n"))
        return paragraphs

    raise ValueError(f"Định dạng không hỗ trợ: {ext}")


def translate_document(input_path: str, output_path: str, target_lang: str) -> str:
    """
    Dịch toàn bộ file và ghi kết quả ra output_path.
    Giữ định dạng gốc: .txt -> .txt, .docx -> .docx, .pdf -> .pdf (dạng văn bản thuần dịch lại).
    """
    ext = _ext(input_path)
    paragraphs = extract_paragraphs(input_path)

    # Dịch theo lô, giữ nguyên các dòng trống để bảo toàn bố cục
    non_empty_idx = [i for i, p in enumerate(paragraphs) if p.strip()]
    texts_to_translate = [paragraphs[i] for i in non_empty_idx]
    translated = translate_batch(
        texts_to_translate, target_lang, context="Đây là nội dung tài liệu, dịch tự nhiên, đúng văn phong."
    )

    result_paragraphs = list(paragraphs)
    for pos, i in enumerate(non_empty_idx):
        result_paragraphs[i] = translated[pos]

    if ext == ".txt":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(result_paragraphs))

    elif ext == ".docx":
        doc = Document()
        for line in result_paragraphs:
            doc.add_paragraph(line)
        doc.save(output_path)

    elif ext == ".pdf":
        styles = getSampleStyleSheet()
        style = styles["Normal"]
        style.alignment = TA_LEFT
        pdf_doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        for line in result_paragraphs:
            if line.strip():
                # escape ký tự đặc biệt XML cơ bản cho reportlab
                safe = (
                    line.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                )
                story.append(Paragraph(safe, style))
        pdf_doc.build(story)

    else:
        raise ValueError(f"Định dạng không hỗ trợ: {ext}")

    return output_path
